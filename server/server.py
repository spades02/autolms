import assemblyai as aai
from flask import Flask, request, jsonify, send_file, make_response
from flask_cors import CORS
from moviepy.editor import *
from werkzeug.utils import secure_filename
import os
import tempfile
import requests
# from openai import OpenAI
import openai
import firebase_admin
from firebase_admin import credentials, storage, auth, db
from docx import Document
from datetime import datetime
import bcrypt
import pyodbc
from io import BytesIO
import requests


connection_string = 'Driver={SQL Server};Server=localhost\SQLEXPRESS;Database=master;Trusted_Connection=True';

# Initialize the database connection
# conn = pyodbc.connect(connection_string)


cred = credentials.Certificate({
  "type": "service_account",
  "project_id": "autolms-a74ce",
  "private_key_id": "cb3c2945439feda17601d5c9c2e75e5b4c82c813",
  "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvgIBADANBgkqhkiG9w0BAQEFAASCBKgwggSkAgEAAoIBAQCk56eHCummWfBB\nz19Cv8e6Ks0WiUba/N8DfWKN51t6T24XgG3h1FXKq5MIMUVqU9tQQZFkOKb3sl1S\nQa7RNWfLfvhb3P9KU9SwSqQ/d1SbzVa6dBWGmOr/M2+GWxmoLOs3Rel6ulXYa6lz\nN/jRFy1esvT6s89Tx0osUsbHMuGn3Uk4RSIGH/VVI292X1hscSMrNjlFYltpjLdT\nF3HQQRU1CzR/Rh5xhRUNnZeSSdtZ6zR4SR1oPmrNUERAEUuHMWl9M9qxAQgr9DS9\nI1gUF8fGv+H1eGZ4a6Hw3k83bRcceEpYr/cBj0vxi8Fbb20OLE2FYe2ms+T2YhHR\nJV3TXbDVAgMBAAECggEAA03E1HbKgZOxtY/TB8JR3Q/4deggLEj+qtI7RD6LXaLz\nwLLJAMpBQcM4OeTK0DuH3E7D++8KoVHIftpkBzr7rl8H2MdPsgyF6VZFfiQE5O4z\nO58Rw1WS97PvGxD+LZlgyQW0jMno6GIXzgAkEwRU213YfKFXmQUYHlU2OE1zqepD\ny/tWti5+HW+0w8hYE2ho0qC8/CfIpPcxkZ8iZ7mqAEx1jm9I2N106TFeeP0xfEIj\nSbLB1YvE+Gaichxmdos4J1JRFmWbwVDOuMH6Pl7TxtRl20UO/KPGKiF2vVv9eKWI\n2RdTK3skP5IrazSbT2pZgRdBeI+6xwaAINRzVjws7QKBgQDYFHE4meAb3wNsglP/\nKhN0AmdDtuwwzpMR4p1DB5NplEDzuA92DxH+ZU6tHKwJCAPJT23BIRC3yGV+QyYJ\nJfaTp1+nKDT5OgAWNUiycMZVz5FjKS8N6O4o/p0bOCUrEqtcGqDIorME3EBsIDk3\ncHGOwsu9JkBinTWC6vvyj9+sbwKBgQDDXuDW/nCI84cc/ARbc8toxEmsRGqzuHKe\n2KahpotokIR7GIxkHV6o9K7ocJzBEwBEyCQtOvj0sSfUdzKB87d7dQfKHA+vOl+8\nvSP07J1C2VERF8H8AcqGawQBuoIOnpBkeCCth/BppkXgLJVzaxpMOZE0ocCnDRSL\n6Pk6Yvpg+wKBgQCgOkBevVO+txdd3iKVIsk8DBAw2TK5WmUaLRV6P+LitmmXkBP+\n8lvgbZTyr0EujT/phjmXrYItFa+U+gv5WqpGgSVf2WO2r3ii/y4CC7g2p21iv4ZB\n5Ui9I1iyd2awUu64pJI3VNY0s3Id+6MR5hJ+zxmvD+9McNDeXIspKeCthwKBgDyP\nfYzHZGTfcGOIr2UV2NjJ97hhoN3C8CzTJZv4P0CTg62Qp3wlKoGfD1TYC49B75Ri\noewvWpqnC5ytskdcH+UyGs8Iscf0hVbBHpM+gWUGbj2pKqWOfJP8Okq64LPdbQ7n\nueFwwRJOy1w5Fb6oVYeJVLyG42wk/th7yD3UDp7fAoGBAMEi1zeFJhVCTBve4Crj\nT+FJ9yNoO6YS1FO9Zw3ZAkp5fyHSx0n+2CX2OXlrNzaBQi6YyHlN9E2vzuW+BZV2\n4W2RovtDAytFzyZFSAkgwoIaR7CjR58JjHnCr+N3jdBrxjQ5lU8mJa4jxuy/I1Z8\n+GGoQZ3qpx47yQgaaPoGcv+c\n-----END PRIVATE KEY-----\n",
  "client_email": "firebase-adminsdk-dyfy3@autolms-a74ce.iam.gserviceaccount.com",
  "client_id": "116793946563957564662",
  "auth_uri": "https://accounts.google.com/o/oauth2/auth",
  "token_uri": "https://oauth2.googleapis.com/token",
  "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
  "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-dyfy3%40autolms-a74ce.iam.gserviceaccount.com",
  "universe_domain": "googleapis.com"
}
)
firebase_admin.initialize_app(cred, {
    'databaseURL': 'https://autolms-a74ce-default-rtdb.asia-southeast1.firebasedatabase.app/'
})

# bucket = storage.bucket()


app = Flask(__name__)
CORS(app)
# client = OpenAI()

def resource_generation(text):

    try:        
        openai.api_key = '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'
        
        print('in function')
        completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": text},
            ]   
        )

        print(completion.choices[0].message)
        sum = completion.choices[0].message.content
        # doc = Document()
        # doc.add_paragraph(sum)  
        # text_file_path = 'summarized_text.docx'
        # doc.save(text_file_path)  
        return sum
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def transcribe_audio(audio_file_url):
    aai.settings.api_key = "f2b25c07671e4d2cac25479469debcab" 

    transcriber = aai.Transcriber()

    audio_url = audio_file_url
    transcript = transcriber.transcribe(audio_url)
    if transcript.error:
        print(transcript.error)
    text = transcript.text
    return text
    # print(transcript.text)

    
    # transcriber = aai.Transcriber()
    # transcript = transcriber.transcribe(audio_file_url)
    # if transcript:
    #     return transcript.text
    # else:
    #     return None

def hash_password(password):
    salt = bcrypt.gensalt()
    hashed_password = bcrypt.hashpw(password.encode('utf-8'), salt)
    return hashed_password

@app.route('/register', methods=['POST'])
def register():
    data = request.json
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')
    sql = "INSERT INTO Users (Username, Email, Password) VALUES (?, ?, ?)"

    print(username, email, password)

    try:
        # cursor = conn.cursor()
        # cursor.execute(sql, (username, email, password))
        # conn.commit()  # Commit the transaction
        # conn.close()
        # print('Success')
        # hashed_password = hash_password(password)
        # print("Hashed Password:", hashed_password)

        return jsonify({'message': 'User registered successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/login', methods=['POST'])
def login():
    data = request.json
    email = data.get('email')
    password = data.get('password')

    try:
        ref = db.reference('user/new_user')
        snapshot = ref.order_by_child('email').equal_to(email).get()
        if snapshot:
            print('user hy')
        else:
            print('user hy nai')

        

    #     # Verify user's credentials using Firebase Authentication
    #     user = auth.get_user_by_email(email)
    #     auth.verify_password(user.uid, password)

    #     # Get additional user data from Firebase Realtime Database
    #     user_data = db.reference('users').child(user.uid).get()

    #     return jsonify({'message': 'User logged in successfully', 'user': user_data}), 200
    # except auth.AuthError as e:
    #     return jsonify({'error': 'Invalid email or password'}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/generate_content', methods=['POST'])
def generate_content():
    try:
        files = []
        total_text_generated = ''
        resources_generated = []
        video_validation = 'None'

        url_count = int(request.form.get('url_count', 0))
        for i in range(url_count):
            url_key = f'url{i}'
            if url_key in request.form:
                url = request.form[url_key]
                response = requests.get(url)
                if response.status_code == 200:
                    file_content = BytesIO(response.content)
                    filename = secure_filename(url.split('/')[-1])  # Set the filename
                    temp_file_path = os.path.join(tempfile.gettempdir(), filename)
                    with open(temp_file_path, 'wb') as f:
                        f.write(file_content.getbuffer())

                    audio_temp_file_path = os.path.join(tempfile.gettempdir(), f'audio{i}.mp3')
                    print('Video File received with the name:', filename)

                    try:
                        video_clip = VideoFileClip(temp_file_path)
                        audio_clip = video_clip.audio
                        audio_clip.write_audiofile(audio_temp_file_path)
                        audio_clip.close()
                        video_clip.close()
                    except Exception as e:
                        print(f"Error processing video {filename}: {e}")
                        return jsonify({'error': 'Error processing video file'}), 500

                    transcription_text = transcribe_audio(audio_temp_file_path)
                    total_text_generated += transcription_text
                    if transcription_text:
                        print(f'Transcription text for audio {i}: {transcription_text}')
                    else:
                        print(f'Failed to transcribe audio {i}')
                        return jsonify({'error': 'Inappropriate Video or Bad Audio'}), 400

        if not total_text_generated:
            return jsonify({'error': 'Inappropriate Video or Bad Audio.'}), 400
        else:
            quiz = request.form.get('quiz')
            asgn = request.form.get('asgn')
            notes = request.form.get('notes')
            project = request.form.get('project')
            paper = request.form.get('paper')
            summary = request.form.get('summary')
            prompt_input = request.form.get('prompt_input')
            language_support = request.form.get('language_support', '')

            video_validation_prompt = f'You are provided with the text that is generated by a video, you are to hereby analyze the text and you are requested to check whether the given video is an educational video or not. Additionally, you are requested to give a one-word answer, yes or no, because your answer will be sent to a backend server and it will expect a yes or no answer from you. So if video is educational, reply yes, else no. Text:{total_text_generated} So if video is educational, reply yes, else no, don’t give any explanation, only give one word answer.'
            video_validation_answer = resource_generation(video_validation_prompt)
            print(video_validation_answer)

            if video_validation_answer.lower() == 'yes':
                if quiz:
                    prompt_gpt = f'Following given content is the text generated from an educational video. Generate a quiz of the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}'
                    resources_generated.append(resource_generation(prompt_gpt))
                if asgn:
                    prompt_gpt = f'Following given content is the text generated from an educational video. Generate an assignment of the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}'
                    resources_generated.append(resource_generation(prompt_gpt))
                if notes:
                    prompt_gpt = f'Following given content is the text generated from an educational video. Generate lecture notes from the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}'
                    resources_generated.append(resource_generation(prompt_gpt))
                if project:
                    prompt_gpt = f'Following given content is the text generated from an educational video. Generate a project idea from the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}'
                    resources_generated.append(resource_generation(prompt_gpt))
                if paper:
                    prompt_gpt = f"Following given content is the text generated from an educational video. Generate a paper with short and long questions from the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}"
                    resources_generated.append(resource_generation(prompt_gpt))
                if summary:
                    prompt_gpt = f'Following given content is the text generated from an educational video. Generate a summary from the following content given. Be restricted to this given text, don’t use your own vector space. Text:{language_support}{total_text_generated}'
                    resources_generated.append(resource_generation(prompt_gpt))
                if prompt_input:
                    prompt_gpt = f'{prompt_input}{total_text_generated} You are hereby requested to generate educational resources for the valid requests from the user.'
                    resources_generated.append(resource_generation(prompt_gpt))

                return jsonify(resources_generated), 200
            else:
                return jsonify({'error': 'Inappropriate video selected! Please upload an educational video.'}), 400

    except Exception as e:
        print(f"Unhandled exception: {e}")
        return jsonify({'error': 'Something went wrong', 'details': str(e)}), 500

@app.route('/download_quiz', methods =['POST'])
def download_quiz():
    try:
        incoming_data = request.get_data(as_text=True)  # Get the incoming data as text
        print("Received data:", incoming_data)

        # Create a new Document
        doc = Document()
        doc.add_paragraph(incoming_data)
        file_path = "quiz.docx"

        # Save the Document to a temporary path
        doc.save(file_path)

        # Get the user's downloads folder
        downloads_folder = os.path.expanduser("~\\Downloads")

        # Define the path where the file will be moved to
        download_path = os.path.join(downloads_folder, "quiz.docx")

        # Check if a file with the same name already exists in the downloads folder
        if os.path.exists(download_path):
            # Generate a unique filename by adding a timestamp to the filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_path = os.path.join(downloads_folder, f"quiz{timestamp}.docx")

        # Move the file to the downloads folder
        os.rename(file_path, download_path)

        # Send the file as an attachment and remove it after sending
        return send_file(download_path, as_attachment=True, attachment_filename="quiz.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', add_etags=False, cache_timeout=0)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_asgn', methods =['POST'])
def download_asgn():
    try:
        incoming_data = request.get_data(as_text=True)  # Get the incoming data as text
        print("Received data:", incoming_data)

        # Create a new Document
        doc = Document()
        doc.add_paragraph(incoming_data)
        file_path = "assignment.docx"

        # Save the Document to a temporary path
        doc.save(file_path)

        # Get the user's downloads folder
        downloads_folder = os.path.expanduser("~\\Downloads")

        # Define the path where the file will be moved to
        download_path = os.path.join(downloads_folder, "assignment.docx")

        # Check if a file with the same name already exists in the downloads folder
        if os.path.exists(download_path):
            # Generate a unique filename by adding a timestamp to the filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_path = os.path.join(downloads_folder, f"assignment_{timestamp}.docx")

        # Move the file to the downloads folder
        os.rename(file_path, download_path)

        # Send the file as an attachment and remove it after sending
        return send_file(download_path, as_attachment=True, attachment_filename="assignment.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', add_etags=False, cache_timeout=0)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_notes', methods =['POST'])
def download_notes():
    try:
        incoming_data = request.get_data(as_text=True)  # Get the incoming data as text
        print("Received data:", incoming_data)

        # Create a new Document
        doc = Document()
        doc.add_paragraph(incoming_data)
        file_path = "notes.docx"

        # Save the Document to a temporary path
        doc.save(file_path)

        # Get the user's downloads folder
        downloads_folder = os.path.expanduser("~\\Downloads")

        # Define the path where the file will be moved to
        download_path = os.path.join(downloads_folder, "notes.docx")

        # Check if a file with the same name already exists in the downloads folder
        if os.path.exists(download_path):
            # Generate a unique filename by adding a timestamp to the filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_path = os.path.join(downloads_folder, f"notes{timestamp}.docx")

        # Move the file to the downloads folder
        os.rename(file_path, download_path)

        # Send the file as an attachment and remove it after sending
        return send_file(download_path, as_attachment=True, attachment_filename="notes.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', add_etags=False, cache_timeout=0)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_proj', methods =['POST'])
def download_proj():
    try:
        incoming_data = request.get_data(as_text=True)  # Get the incoming data as text
        print("Received data:", incoming_data)

        # Create a new Document
        doc = Document()
        doc.add_paragraph(incoming_data)
        file_path = "project_idea.docx"

        # Save the Document to a temporary path
        doc.save(file_path)

        # Get the user's downloads folder
        downloads_folder = os.path.expanduser("~\\Downloads")

        # Define the path where the file will be moved to
        download_path = os.path.join(downloads_folder, "project_idea.docx")

        # Check if a file with the same name already exists in the downloads folder
        if os.path.exists(download_path):
            # Generate a unique filename by adding a timestamp to the filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_path = os.path.join(downloads_folder, f"project_idea{timestamp}.docx")

        # Move the file to the downloads folder
        os.rename(file_path, download_path)

        # Send the file as an attachment and remove it after sending
        return send_file(download_path, as_attachment=True, attachment_filename="project_idea.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', add_etags=False, cache_timeout=0)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_additionals', methods =['POST'])
def download_additionals():
    try:
        incoming_data = request.get_data(as_text=True)  # Get the incoming data as text
        print("Received data:", incoming_data)

        # Create a new Document
        doc = Document()
        doc.add_paragraph(incoming_data)
        file_path = "additional_resources.docx"

        # Save the Document to a temporary path
        doc.save(file_path)

        # Get the user's downloads folder
        downloads_folder = os.path.expanduser("~\\Downloads")

        # Define the path where the file will be moved to
        download_path = os.path.join(downloads_folder, "additional_resources.docx")

        # Check if a file with the same name already exists in the downloads folder
        if os.path.exists(download_path):
            # Generate a unique filename by adding a timestamp to the filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_path = os.path.join(downloads_folder, f"additional_resources{timestamp}.docx")

        # Move the file to the downloads folder
        os.rename(file_path, download_path)

        # Send the file as an attachment and remove it after sending
        return send_file(download_path, as_attachment=True, attachment_filename="additional_resources.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', add_etags=False, cache_timeout=0)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000)
